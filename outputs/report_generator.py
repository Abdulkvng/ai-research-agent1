
from docx import Document

def generate_doc(summary, q_and_a):
    doc = Document()
    doc.add_heading("AI Research Assistant Output", 0)
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)
    doc.add_heading("Q&A", level=1)
    for q, a in q_and_a.items():
        doc.add_paragraph(f"Q: {q}")
        doc.add_paragraph(f"A: {a}")
    doc.save("output.docx")
